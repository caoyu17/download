# 使用 pip 安装依赖：pip install flask python-docx cos-python-sdk-v5
from flask import Flask, request, jsonify
from docx import Document
from io import BytesIO
from qcloud_cos import CosConfig
from qcloud_cos import CosS3Client
import logging
import sys

# 配置日志
logging.basicConfig(level=logging.INFO, stream=sys.stdout)

# Flask 应用
app = Flask(__name__)

# 腾讯云 COS 配置（请替换为实际值）
COS_SECRET_ID = ''
COS_SECRET_KEY = ''
COS_REGION = 'ap-guangzhou'  # 替换为存储桶所属地域
COS_BUCKET = 'xiaodao-1331856197'

# 初始化 COS 客户端
config = CosConfig(Region=COS_REGION, SecretId=COS_SECRET_ID, SecretKey=COS_SECRET_KEY)
cos_client = CosS3Client(config)


def markdown_to_docx(markdown_text):
    """将 Markdown 文本转换为 DOCX 文档"""
    doc = Document()
    lines = markdown_text.split('\n')

    for line in lines:
        if line.startswith('#'):
            level = min(line.count('#'), 9)  # 限制标题级别最大为 9
            doc.add_heading(line.lstrip('#').strip(), level)
        elif line.startswith('* '):
            doc.add_paragraph(line.lstrip('* '), style='List Bullet')
        elif line.startswith('1. '):
            doc.add_paragraph(line.lstrip('1. '), style='List Number')
        else:
            doc.add_paragraph(line)

    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io


def upload_to_cos(file_obj, cos_key):
    """上传文件对象到 COS"""
    try:
        response = cos_client.put_object(
            Bucket=COS_BUCKET,
            Body=file_obj,
            Key=cos_key,
            ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'  # DOCX 标准 MIME 类型
        )
        logging.info(f"上传成功至 COS: {cos_key}, ETag: {response['ETag']}")
        return True, cos_key  # 同时返回 cos_key，方便后续使用
    except Exception as e:
        logging.error(f"上传失败: {e}")
        return False, None


@app.route('/convert-and-upload', methods=['POST'])
def convert_and_upload():
    """处理 Markdown 文本，转换为 DOCX 并上传至 COS"""
    # 获取请求中的 Markdown 文本
    markdown_text = request.form.get('markdown_text')
    if not markdown_text:
        return jsonify({"error": "请提供 Markdown 文本"}), 400

    # 生成唯一文件名（使用时间戳避免冲突）
    import time
    timestamp = int(time.time())
    cos_key = f"docx_files/converted_{timestamp}.docx"

    try:
        # 转换 Markdown 为 DOCX
        docx_io = markdown_to_docx(markdown_text)

        # 上传到 COS
        upload_success = upload_to_cos(docx_io, cos_key)
        if not upload_success:
            return jsonify({"error": "上传至 COS 失败"}), 500

        # （可选）下载到本地路径（如果需要）
        # download_success = download_from_cos(cos_key, f"./downloaded_{timestamp}.docx")
        # if not download_success:
        #     return jsonify({"error": "下载文件失败"}), 500

        # 返回成功响应，包含 COS 中的文件路径
        return jsonify({
            "message": "转换并上传成功",
            "cos_key": cos_key,
            "download_url": f"https://{COS_BUCKET}.cos.{COS_REGION}.myqcloud.com/{cos_key}"
        }), 200

    except Exception as e:
        logging.error(f"处理请求失败: {e}")
        return jsonify({"error": f"处理请求失败: {str(e)}"}), 500


import tempfile  # 用于创建临时文件


@app.route('/download-from-cos/<path:cos_key>', methods=['GET'])
def download_from_cos_api(cos_key):
    try:
        # 1. 生成一个临时文件路径（确保是字符串类型，符合系统路径规则）
        import tempfile
        import os
        # 创建临时文件，返回路径（如 C:\Users\XXX\AppData\Local\Temp\tmpXXXXXX.docx）
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
            temp_path = temp_file.name  # 这是一个有效的字符串路径

        # 2. 调用 download_file 时，传入临时文件路径（字符串类型）
        cos_client.download_file(
            Bucket=COS_BUCKET,
            Key=cos_key,
            DestFilePath=temp_path  # 正确：传入字符串路径
        )

        # 3. 读取临时文件内容，返回给客户端
        with open(temp_path, 'rb') as f:
            file_stream = f.read()

        # 4. 清理临时文件（避免占用磁盘）
        os.unlink(temp_path)

        # 5. 构造响应返回
        from flask import make_response
        resp = make_response(file_stream)
        resp.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        filename = cos_key.split("/")[-1]
        resp.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        resp.headers['Content-Length'] = str(len(file_stream))

        return resp
    except Exception as e:
        logging.error(f"接口下载失败: {e}")
        return jsonify({"error": "下载文件失败"}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
