from flask import Flask, request, send_file
from docx import Document
from io import BytesIO


app = Flask(__name__)


def markdown_to_docx(markdown_text):
    doc = Document()
    lines = markdown_text.split('\n')
    for line in lines:
        if line.startswith('#'):
            level = line.count('#')
            doc.add_heading(line.lstrip('#').strip(), level)
        elif line.startswith('* '):
            doc.add_paragraph(line.lstrip('* '), style='List Bullet')
        elif line.startswith('1. '):
            doc.add_paragraph(line.lstrip('1. '), style='List Number')
        else:
            doc.add_paragraph(line)
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io


@app.route('/download', methods=['POST'])
def download():
    markdown_text = request.form.get('markdown_text')
    if not markdown_text:
        return '请提供Markdown文本', 400
    docx_io = markdown_to_docx(markdown_text)
    return send_file(
        docx_io,
        as_attachment=True,
        download_name='converted_docx.docx',
        mimetype='application/vnd.openxmlformats - officedocument.wordprocessingml.document'
    )


if __name__ == '__main__':
    app.run(debug=True,host='0.0.0.0')
